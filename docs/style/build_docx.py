"""
build_docx.py -- render a Rayquaza markdown document to an EDITABLE Word
(.docx) file that carries the same design system as the PDF (build_pdf.py).

A Word document is laid out by Word's own engine and cannot be a pixel clone
of the Chromium-rendered PDF (line and page breaks differ), but every design
token is ported faithfully: the paper background, Public Sans / Roboto Mono
(embedded into the file so they render on any machine, even without the fonts
installed), blue section numbers and figure/table numbers, surface-shaded
table headers, shaded inline code, FINDING/HEADLINE pills, the cover page, and
a running header/footer with a real Word page-number field.

Content comes from the SAME markdown->HTML pipeline used for the PDF
(md_to_html.render_markdown_to_html); this module parses that known,
well-structured HTML and emits the matching Word elements.

Run: python docs/style/build_docx.py docs/paper/paper.md
"""
import argparse
import base64
import io
import uuid
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from md_to_html import render_markdown_to_html

STYLE_DIR = Path(__file__).resolve().parent
FONTS_DIR = STYLE_DIR / "fonts"

# --- design tokens (mirror of tokens.css) ---
PAPER = "F9F8F3"
SURFACE = "E2E1DA"
BORDER = "BEBEBE"
INK = "262626"
INK_SOFT = "313131"
BLUE = "0099FF"
GREEN = "2FBB45"
ORANGE = "DC762D"
RED = "FB2C55"

FONT_SANS = "Public Sans"
FONT_MONO = "Roboto Mono"

BODY_PT = 10.5
BODY_LINE = 1.35


# ======================================================================
# HTML parsing -- build a tiny DOM from the (known-shape) body HTML
# ======================================================================
class _Node:
    __slots__ = ("tag", "attrs", "children", "text")

    def __init__(self, tag, attrs=None):
        self.tag = tag
        self.attrs = dict(attrs or {})
        self.children = []
        self.text = None  # for text nodes, tag == "#text"


class _DOMParser(HTMLParser):
    VOID = {"img", "br", "hr", "meta", "link", "input"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _Node("#root")
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        node = _Node(tag, attrs)
        self.stack[-1].children.append(node)
        if tag not in self.VOID:
            self.stack.append(node)

    def handle_startendtag(self, tag, attrs):
        self.stack[-1].children.append(_Node(tag, attrs))

    def handle_endtag(self, tag):
        if tag in self.VOID:
            return
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                break

    def handle_data(self, data):
        n = _Node("#text")
        n.text = data
        self.stack[-1].children.append(n)


def _parse_html(html: str) -> _Node:
    p = _DOMParser()
    p.feed(html)
    return p.root


def _cls(node: _Node) -> list[str]:
    return node.attrs.get("class", "").split()


# ======================================================================
# low-level docx helpers
# ======================================================================
def _set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_margins(cell, top=40, bottom=40, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _set_table_width_pct(table, pct=5000):
    """Set a table's preferred width to a percentage of the text area
    (5000 = 100%), so wide tables use the full page width and their columns
    get enough room to avoid brutal mid-token wrapping."""
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(pct))
    tblW.set(qn("w:type"), "pct")


def _set_table_borders(table, color=BORDER, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def _shade_run(run, hex_fill):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    rPr.append(shd)


def _style_run(run, *, font=FONT_SANS, size=BODY_PT, color=INK,
               bold=False, italic=False, mono=False, shade=None):
    run.font.name = FONT_MONO if mono else font
    # ensure east-asian/cs also use the font so Word doesn't substitute
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    name = FONT_MONO if mono else font
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(a), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic
    if shade:
        _shade_run(run, shade)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldStart = OxmlElement("w:fldChar"); fldStart.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldStart); run._r.append(instr); run._r.append(fldEnd)
    _style_run(run, mono=True, size=8, color=INK_SOFT)


# ======================================================================
# inline runs: walk an element's inline children (#text, strong, em, code, a)
# ======================================================================
def _emit_inline(paragraph, node, *, base_size=BODY_PT, base_color=INK,
                 bold=False, italic=False):
    for child in node.children:
        if child.tag == "#text":
            if child.text:
                r = paragraph.add_run(child.text)
                _style_run(r, size=base_size, color=base_color, bold=bold, italic=italic)
        elif child.tag in ("strong", "b"):
            _emit_inline(paragraph, child, base_size=base_size, base_color=base_color,
                         bold=True, italic=italic)
        elif child.tag in ("em", "i"):
            _emit_inline(paragraph, child, base_size=base_size, base_color=base_color,
                         bold=bold, italic=True)
        elif child.tag == "code":
            txt = _text_of(child)
            r = paragraph.add_run(txt)
            _style_run(r, mono=True, size=base_size - 1.5, color=INK, shade=SURFACE)
        elif child.tag == "a":
            txt = _text_of(child)
            r = paragraph.add_run(txt)
            _style_run(r, size=base_size, color=BLUE)
        elif child.tag == "span" and "secnum" in _cls(child):
            r = paragraph.add_run(_text_of(child) + " ")
            _style_run(r, size=base_size, color=BLUE, bold=True)
        elif child.tag == "span" and "fignum" in _cls(child):
            r = paragraph.add_run(_text_of(child) + " ")
            _style_run(r, mono=True, size=base_size, color=BLUE, bold=True)
        elif child.tag == "span" and "reftag" in _cls(child):
            r = paragraph.add_run(_text_of(child))
            _style_run(r, mono=True, size=8, color=BLUE, shade=SURFACE)
            paragraph.add_run(" ")
        elif child.tag == "br":
            paragraph.add_run().add_break()
        else:
            # unknown inline wrapper: recurse
            _emit_inline(paragraph, child, base_size=base_size, base_color=base_color,
                         bold=bold, italic=italic)


def _text_of(node: _Node) -> str:
    if node.tag == "#text":
        return node.text or ""
    return "".join(_text_of(c) for c in node.children)


# ======================================================================
# block emitters
# ======================================================================
def _emit_heading(doc, node, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    size = 15 if level == 2 else 12.5
    # secnum span (blue) + remaining text
    for child in node.children:
        if child.tag == "span" and "secnum" in _cls(child):
            r = p.add_run(_text_of(child) + "  ")
            _style_run(r, size=size, color=BLUE, bold=True)
        elif child.tag == "#text":
            if child.text and child.text.strip():
                r = p.add_run(child.text)
                _style_run(r, size=size, color=INK, bold=True)
        else:
            r = p.add_run(_text_of(child))
            _style_run(r, size=size, color=INK, bold=True)


def _emit_paragraph(doc, node):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = BODY_LINE
    _emit_inline(p, node)
    return p


def _emit_list(doc, node, ordered):
    idx = 1
    for li in node.children:
        if li.tag != "li":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = BODY_LINE
        marker = f"{idx}. " if ordered else "•  "
        r = p.add_run(marker)
        _style_run(r, size=BODY_PT, color=INK, bold=ordered)
        _emit_inline(p, li)
        idx += 1


def _emit_callout(doc, node):
    kind = "finding" if "callout--finding" in _cls(node) else \
           ("headline" if "callout--headline" in _cls(node) else "")
    pill_bg = RED if kind == "finding" else (ORANGE if kind == "headline" else INK)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    _set_table_width_pct(tbl)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, SURFACE)
    _set_cell_margins(cell, top=90, bottom=90, left=140, right=140)
    # remove default paragraph, build our own
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = BODY_LINE
    # find the pill span + remaining inline content
    pill = None
    rest_children = []
    for child in node.children:
        if child.tag == "span" and "pill" in _cls(child):
            pill = _text_of(child)
        else:
            rest_children.append(child)
    if pill:
        r = p.add_run(f" {pill} ")
        _style_run(r, mono=True, size=8, color=PAPER, bold=True, shade=pill_bg)
        p.add_run("  ")
    holder = _Node("holder"); holder.children = rest_children
    _emit_inline(p, holder)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _emit_table(doc, table_node):
    # locate thead/tbody rows
    thead = _find(table_node, "thead")
    tbody = _find(table_node, "tbody")
    header_cells = []
    if thead:
        tr = _find(thead, "tr")
        header_cells = [c for c in tr.children if c.tag in ("th", "td")]
    body_rows = []
    if tbody:
        for tr in tbody.children:
            if tr.tag == "tr":
                body_rows.append([c for c in tr.children if c.tag in ("td", "th")])
    ncols = len(header_cells) or (len(body_rows[0]) if body_rows else 1)

    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.autofit = True
    _set_table_borders(tbl)
    _set_table_width_pct(tbl)

    if header_cells:
        row = tbl.add_row()
        for i, hc in enumerate(header_cells):
            cell = row.cells[i]
            _set_cell_bg(cell, SURFACE)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].text = ""
            _emit_cell_inline(cell.paragraphs[0], hc, header=True)
    for br in body_rows:
        row = tbl.add_row()
        for i in range(ncols):
            cell = row.cells[i]
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].text = ""
            if i < len(br):
                _emit_cell_inline(cell.paragraphs[0], br[i], header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _emit_cell_inline(paragraph, cell_node, header):
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.space_after = Pt(0)
    size = 8.5
    for child in cell_node.children:
        if child.tag == "#text":
            if child.text:
                r = paragraph.add_run(child.text)
                _style_run(r, size=size, color=INK, bold=header)
        elif child.tag == "code":
            r = paragraph.add_run(_text_of(child))
            _style_run(r, mono=True, size=size - 0.5, color=INK, shade=SURFACE)
        elif child.tag in ("strong", "b"):
            r = paragraph.add_run(_text_of(child))
            _style_run(r, size=size, color=INK, bold=True)
        elif child.tag in ("em", "i"):
            r = paragraph.add_run(_text_of(child))
            _style_run(r, size=size, color=INK, italic=True)
        else:
            r = paragraph.add_run(_text_of(child))
            _style_run(r, size=size, color=INK, bold=header)


def _emit_figure(doc, fig_node):
    img = _find(fig_node, "img")
    cap = _find(fig_node, "figcaption")
    if img is not None:
        src = img.attrs.get("src", "")
        if src.startswith("data:image/png;base64,"):
            data = base64.b64decode(src.split(",", 1)[1])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run()
            run.add_picture(io.BytesIO(data), width=Inches(6.3))
    if cap is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for child in cap.children:
            if child.tag == "span" and "fignum" in _cls(child):
                r = p.add_run(_text_of(child) + " ")
                _style_run(r, mono=True, size=8, color=BLUE, bold=True)
            elif child.tag == "#text":
                if child.text:
                    r = p.add_run(child.text)
                    _style_run(r, mono=True, size=8, color=INK_SOFT)
            else:
                r = p.add_run(_text_of(child))
                _style_run(r, mono=True, size=8, color=INK_SOFT)


def _emit_table_caption(doc, node):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    for child in node.children:
        if child.tag == "span" and "fignum" in _cls(child):
            r = p.add_run(_text_of(child) + " ")
            _style_run(r, mono=True, size=8, color=BLUE, bold=True)
        elif child.tag == "#text":
            if child.text:
                r = p.add_run(child.text)
                _style_run(r, mono=True, size=8, color=INK_SOFT)
        else:
            r = p.add_run(_text_of(child))
            _style_run(r, mono=True, size=8, color=INK_SOFT)


def _emit_pre(doc, pre_node):
    code = _find(pre_node, "code") or pre_node
    text = _text_of(code)
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_width_pct(tbl)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, SURFACE)
    _set_cell_margins(cell, top=80, bottom=80, left=110, right=110)
    _set_table_borders(tbl)
    cell.paragraphs[0].text = ""
    for j, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _style_run(r, mono=True, size=8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _find(node: _Node, tag: str):
    for c in node.children:
        if c.tag == tag:
            return c
    for c in node.children:
        found = _find(c, tag)
        if found is not None:
            return found
    return None


# ======================================================================
# document assembly
# ======================================================================
def _walk_body(doc, root: _Node):
    # body content lives directly under #root
    for node in root.children:
        t = node.tag
        if t == "h2":
            _emit_heading(doc, node, 2)
        elif t == "h3":
            _emit_heading(doc, node, 3)
        elif t == "p":
            if "table-caption" in _cls(node):
                _emit_table_caption(doc, node)
            else:
                _emit_paragraph(doc, node)
        elif t == "ul":
            _emit_list(doc, node, ordered=False)
        elif t == "ol":
            _emit_list(doc, node, ordered=True)
        elif t == "div" and "callout" in _cls(node):
            _emit_callout(doc, node)
        elif t == "div" and "table-wrap" in _cls(node):
            table = _find(node, "table")
            if table is not None:
                _emit_table(doc, table)
        elif t == "table":
            _emit_table(doc, node)
        elif t == "figure":
            _emit_figure(doc, node)
        elif t == "pre":
            _emit_pre(doc, node)
        elif t == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), BORDER)
            pbdr.append(bottom); pPr.append(pbdr)


def _set_page_background(doc, hex_color):
    """Set the Word page background colour (document background + the display
    flag in settings)."""
    # <w:background w:color="..."/> as first child of w:document
    body = doc.element.body
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    doc.element.insert(0, bg)
    # settings: displayBackgroundShape
    settings = doc.settings.element
    disp = OxmlElement("w:displayBackgroundShape")
    settings.append(disp)


def _build_cover(doc, fm):
    title = fm.get("title", "")
    authors = fm.get("authors", "")
    affil = fm.get("affiliation", "")
    date = fm.get("date", "")
    cls = fm.get("classification", "")

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(affil.upper())
    _style_run(r, mono=True, size=8.5, color=INK_SOFT)
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    r = p.add_run(title)
    _style_run(r, size=23, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(20)
    p.paragraph_format.line_spacing = 1.15

    p = doc.add_paragraph()
    r = p.add_run(authors)
    _style_run(r, size=12.5, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run(affil)
    _style_run(r, size=10.5, color=INK_SOFT)
    p.paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph()
    r = p.add_run(f"{date}    ·    {cls}")
    _style_run(r, mono=True, size=9, color=INK_SOFT)


def _configure_running_header_footer(section, short_title, classification):
    # header: short title (mono, ink-soft)
    hp = section.header.paragraphs[0]
    hp.text = ""
    r = hp.add_run(short_title)
    _style_run(r, mono=True, size=8, color=INK_SOFT)
    # footer: classification (left) + page number (right) via a tab
    fp = section.footer.paragraphs[0]
    fp.text = ""
    r = fp.add_run(classification)
    _style_run(r, mono=True, size=8, color=INK_SOFT)
    # right-aligned tab stop for the page number
    from docx.enum.text import WD_TAB_ALIGNMENT
    tabs = fp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.9), WD_TAB_ALIGNMENT.RIGHT)
    fp.add_run("\t")
    _add_page_number_field(fp)


# ---- font embedding (OOXML obfuscated fonts) ----
def _obfuscate_font(font_bytes: bytes, guid: str) -> bytes:
    """Obfuscate a font per ECMA-376: XOR the first 32 bytes with the 16-byte
    GUID (as hex, reversed), applied twice."""
    key = bytes.fromhex(guid.replace("-", ""))[::-1]
    data = bytearray(font_bytes)
    for i in range(32):
        data[i] ^= key[i % 16]
    return bytes(data)


def _embed_fonts_in_zip(docx_path: Path) -> None:
    """Embed Public Sans (regular/italic/bold) + Roboto Mono (regular) into an
    already-saved .docx by rewriting the package zip. Done as a zip
    post-process (rather than through python-docx's OPC model) to avoid
    duplicating the default fontTable part. Word and LibreOffice both read
    these embedded, obfuscated fonts, so the document renders correctly even
    where the fonts are not installed."""
    import re as _re
    import zipfile

    faces = {
        FONT_SANS: [
            ("embedRegular", FONTS_DIR / "publicsans-400.ttf"),
            ("embedBold", FONTS_DIR / "publicsans-600.ttf"),
            ("embedItalic", FONTS_DIR / "publicsans-400i.ttf"),
        ],
        FONT_MONO: [
            ("embedRegular", FONTS_DIR / "robotomono-400.ttf"),
        ],
    }
    FONT_CT = "application/vnd.openxmlformats-officedocument.obfuscatedFont"
    FONT_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    # build fontTable.xml, its .rels, and the obfuscated font blobs
    fonts_xml = [f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                 f'<w:fonts xmlns:w="{W}" xmlns:r="{R}">']
    rels = [f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<Relationships xmlns="{R}">']
    font_blobs = {}  # partname (in zip) -> bytes
    idx = 0
    for family, styles in faces.items():
        fonts_xml.append(f'<w:font w:name="{family}">'
                         f'<w:charset w:val="00"/><w:family w:val="auto"/>'
                         f'<w:pitch w:val="variable"/>')
        for tag, path in styles:
            if not path.exists():
                continue
            idx += 1
            guid = str(uuid.uuid4()).upper()
            obf = _obfuscate_font(path.read_bytes(), guid)
            rid = f"rIdFont{idx}"
            fname = f"font{idx}.odttf"
            font_blobs[f"word/fonts/{fname}"] = obf
            rels.append(f'<Relationship Id="{rid}" Type="{FONT_RT}" Target="fonts/{fname}"/>')
            fonts_xml.append(f'<w:{tag} r:id="{rid}" w:fontKey="{{{guid}}}" w:subsetted="false"/>')
        fonts_xml.append('</w:font>')
    fonts_xml.append('</w:fonts>')
    rels.append('</Relationships>')
    fonts_xml_bytes = "".join(fonts_xml).encode("utf-8")
    rels_bytes = "".join(rels).encode("utf-8")

    # read the whole package, then rewrite with modifications
    src = zipfile.ZipFile(docx_path, "r")
    items = {info.filename: src.read(info.filename) for info in src.infolist()}
    src.close()

    # [Content_Types].xml : add odttf default + ensure fontTable override
    ct = items["[Content_Types].xml"].decode("utf-8")
    if 'Extension="odttf"' not in ct:
        ct = ct.replace("</Types>",
                        f'<Default Extension="odttf" ContentType="{FONT_CT}"/></Types>')
    if "/word/fontTable.xml" not in ct:
        ct = ct.replace("</Types>",
                        '<Override PartName="/word/fontTable.xml" '
                        'ContentType="application/vnd.openxmlformats-officedocument.'
                        'wordprocessingml.fontTable+xml"/></Types>')
    items["[Content_Types].xml"] = ct.encode("utf-8")

    # document.xml.rels : ensure a relationship to fontTable.xml
    doc_rels_name = "word/_rels/document.xml.rels"
    doc_rels = items[doc_rels_name].decode("utf-8")
    if "fontTable.xml" not in doc_rels:
        ft_rt = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable"
        doc_rels = doc_rels.replace(
            "</Relationships>",
            f'<Relationship Id="rIdFontTable" Type="{ft_rt}" Target="fontTable.xml"/></Relationships>')
        items[doc_rels_name] = doc_rels.encode("utf-8")

    # settings.xml : add embedTrueTypeFonts flag (must sit in the correct
    # schema-ordered position -- right after the opening <w:settings ...>)
    settings_name = "word/settings.xml"
    settings = items[settings_name].decode("utf-8")
    if "embedTrueTypeFonts" not in settings:
        settings = _re.sub(r"(<w:settings\b[^>]*>)",
                           r"\1<w:embedTrueTypeFonts/>", settings, count=1)
        items[settings_name] = settings.encode("utf-8")

    # replace fontTable.xml, add its rels + the font blobs
    items["word/fontTable.xml"] = fonts_xml_bytes
    items["word/_rels/fontTable.xml.rels"] = rels_bytes
    items.update(font_blobs)

    # write the new package
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as out:
        for name, data in items.items():
            out.writestr(name, data)


def build_docx(md_path: Path, output_path: Path) -> None:
    front_matter, body_html = render_markdown_to_html(md_path)
    root = _parse_html(body_html)

    doc = Document()

    # base Normal style
    normal = doc.styles["Normal"]
    normal.font.name = FONT_SANS
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor.from_string(INK)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT_SANS)

    # page geometry (A4, margins matching the PDF's spirit)
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.different_first_page_header_footer = True  # cover has no running h/f

    _set_page_background(doc, PAPER)

    # cover on page 1
    _build_cover(doc, front_matter)

    # start the body on a new section (own header/footer + page numbering)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Inches(8.27)
    body_section.page_height = Inches(11.69)
    body_section.left_margin = Inches(0.85)
    body_section.right_margin = Inches(0.85)
    body_section.top_margin = Inches(0.9)
    body_section.bottom_margin = Inches(0.9)
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    # restart page numbering at 1 for the body? keep continuous; cover is p1.
    short_title = str(front_matter.get("title", ""))[:70]
    classification = str(front_matter.get("classification", ""))
    _configure_running_header_footer(body_section, short_title, classification)

    _walk_body(doc, root)

    doc.save(str(output_path))
    _embed_fonts_in_zip(Path(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Render a Rayquaza markdown document to a styled .docx.")
    parser.add_argument("source", type=Path)
    parser.add_argument("--output", type=Path, default=None)
    args = parser.parse_args()
    output = args.output or args.source.with_suffix(".docx")
    build_docx(args.source, output)
    print(f"wrote {output}")


if __name__ == "__main__":
    main()
